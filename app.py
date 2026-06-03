import re
import textwrap
import os
import markdown
import sqlite3
from io import BytesIO
from flask import Flask, render_template, request, send_file
from google import genai
from dotenv import load_dotenv
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document

app = Flask(__name__)

# Load environment variables from .env file
load_dotenv()
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))


def init_db():
    """Creates the dokumentationen table if it does not exist yet."""
    conn = sqlite3.connect("dokumentationen.db")
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS dokumentationen (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            inhalt TEXT NOT NULL
        )
    """)
    conn.commit()
    conn.close()


def get_db():
    """Opens a SQLite connection and returns the connection and cursor."""
    conn = sqlite3.connect("dokumentationen.db")
    cursor = conn.cursor()
    return conn, cursor


def clean_for_html(content):
    """Converts Markdown text to HTML including table support."""
    return markdown.markdown(content, extensions=["tables"])


@app.route("/")
def index():
    """Renders the main page."""
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():
    """
    Receives a process name from the form, sends it to Gemini AI
    and returns a structured process documentation.
    The result is saved to the database and rendered as HTML.
    """
    process = request.form.get("prozess_name")

    if not process:
        return render_template("index.html", error="Es muss etwas eingegeben werden")

    # Build the prompt with a fixed structure for the AI output
    prompt = f"""
       Erstelle eine strukturierte Prozessdokumentation für: {process}

       Bitte gliedere sie so:
       1. Ziel des Prozesses
       2. Schritte (nummeriert)
       3. Verantwortliche pro Schritt
       4. 4. Checkliste am Ende (als normale nummerierte Liste, keine Checkboxen, keine Klammern)
       """

    response = client.models.generate_content(
        model="gemini-3.1-flash-lite",
        contents=prompt
    )

    result_text = response.text

    # Save the raw Markdown text to the database
    conn, cursor = get_db()
    cursor.execute("INSERT INTO dokumentationen (inhalt) VALUES (?)", (result_text,))
    conn.commit()
    conn.close()

    # Convert Markdown to HTML for display
    result_html = clean_for_html(response.text)
    return render_template("index.html", result=result_html, result_text=result_text)


@app.route("/pdf_download", methods=["POST"])
def file_download():
    """
    Generates a PDF from the documentation text.
    Strips Markdown syntax and handles line wrapping and page breaks.
    """
    content = request.form.get("content")
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)

    # Normalize line endings and strip Markdown syntax for plain text output
    content = content.replace('\r\n', '\n').replace('\r', '\n')
    content = re.sub(r'#{1,6}\s*', '', content)
    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
    content = re.sub(r'\*', '', content)

    width, height = A4
    y = height - 50

    for line in content.split("\n"):
        # Start a new page if there is not enough space
        if y < 50:
            p.showPage()
            y = height - 50
        # Wrap long lines so they fit within the page width
        for row in textwrap.wrap(line, width=90):
            if y < 50:
                p.showPage()
                y = height - 50
            p.drawString(50, y, row)
            y -= 20

    p.save()
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype="application/pdf",
        as_attachment=True,
        download_name="dokumentation.pdf"
    )


@app.route("/docx_download", methods=["POST"])
def docx_download():
    """
    Generates a DOCX file from the documentation text.
    Converts Markdown headings to Word headings and strips remaining symbols.
    """
    content = request.form.get("content")
    doc = Document()

    # Normalize line endings
    content = content.replace('\r\n', '\n').replace('\r', '\n')

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Map Markdown headings to Word heading levels
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            # Strip bold markers and add as normal paragraph
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            line = re.sub(r'\*', '', line)
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="dokumentation.docx"
    )


@app.route("/history")
def verlauf():
    """
    Loads all saved documentations from the database,
    converts them from Markdown to HTML and renders the history page.
    """
    conn, cursor = get_db()
    cursor.execute("SELECT * FROM dokumentationen")
    entries = cursor.fetchall()
    conn.close()
    # Convert each entry's Markdown content to HTML before passing to template
    entries = [(e[0], clean_for_html(e[1])) for e in entries]
    return render_template("history.html", entries=entries)


if __name__ == "__main__":
    init_db()
    app.run(debug=True)
